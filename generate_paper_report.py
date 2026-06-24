from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("实验三报告_手写数字视频识别.docx")


def set_font(run, cn="宋体", en="Times New Roman", size=12, bold=False, color=None):
    run.font.name = en
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), cn)
    r_fonts.set(qn("w:ascii"), en)
    r_fonts.set(qn("w:hAnsi"), en)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    set_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_p(
    doc,
    text="",
    size=12,
    bold=False,
    align=None,
    first=False,
    after=0,
    before=0,
    line=2.0,
    cn="宋体",
):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    para.paragraph_format.line_spacing = line
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.space_before = Pt(before)
    if first:
        para.paragraph_format.first_line_indent = Cm(0.74)
    if text:
        run = para.add_run(text)
        set_font(run, cn=cn, size=size, bold=bold)
    return para


def heading(doc, text):
    return add_p(doc, text, size=12, bold=True, before=6, line=2.0)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, size=10.5, bold=True)
        shade(cell, "D9EAF7")
        if widths:
            cell.width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), size=10.5)
            if widths:
                cells[i].width = Cm(widths[i])
    add_p(doc, "", line=1.0)
    return table


def add_placeholder(doc, title, note):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, "F2F2F2")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(f"{title}\n{note}")
    set_font(run, size=12, bold=True, color=(90, 90, 90))
    add_p(
        doc,
        title.replace("【", "").replace("】", ""),
        size=10.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line=1.5,
    )


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    add_p(
        doc,
        "研 究 生 课 程 论 文",
        size=22,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line=1.5,
        cn="黑体",
    )
    add_p(
        doc,
        "(2025-2026学年第二学期)",
        size=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line=1.5,
    )
    for _ in range(2):
        add_p(doc, "", line=1.0)

    cover = doc.add_table(rows=5, cols=4)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover.style = "Table Grid"
    cover_data = [
        ["学    号", "（请填写）", "学    院", "信息与通信工程学院"],
        ["课程编号", "（请填写）", "课程名称", "模式识别与机器视觉"],
        ["学位类别", "硕士", "任课教师", "（请填写）"],
    ]
    for row_idx, row in enumerate(cover_data):
        for col_idx, value in enumerate(row):
            set_cell_text(cover.cell(row_idx, col_idx), value, bold=(col_idx in (0, 2)))
    cover.rows[3].cells[0].merge(cover.rows[3].cells[3])
    set_cell_text(cover.cell(3, 0), "教师评语：\n\n\n\n\n", align=WD_ALIGN_PARAGRAPH.LEFT)
    cover.rows[4].cells[0].merge(cover.rows[4].cells[3])
    set_cell_text(
        cover.cell(4, 0),
        "成绩评定：        分          任课教师签名：                    年    月    日",
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )

    for _ in range(4):
        add_p(doc, "", line=1.0)
    add_p(
        doc,
        "基于视频信息处理的手写数字识别系统设计与实现",
        size=18,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line=1.5,
        cn="黑体",
    )
    add_p(
        doc,
        "研究生：（请填写姓名）",
        size=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line=1.5,
        cn="仿宋_GB2312",
    )
    add_p(
        doc,
        "提交日期： 2026年6月24日                    研究生签名：",
        size=12,
        bold=True,
        line=1.5,
    )
    doc.add_page_break()

    add_p(
        doc,
        "基于视频信息处理的手写数字识别系统设计与实现",
        size=18,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line=2.0,
    )
    add_p(doc, "（请填写姓名）", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    para = add_p(doc, "", line=2.0)
    para.paragraph_format.first_line_indent = Cm(0.74)
    run = para.add_run("摘要：")
    set_font(run, bold=True)
    run = para.add_run(
        "本文围绕实验三“基于视频信息处理的手写数字识别”任务，设计并实现了一个完整的手写数字视频实时识别系统。"
        "系统采用 Python 语言开发，基于 OpenCV 完成视频读取、图像预处理、轮廓检测和 KNN 分类，基于 PyTorch "
        "实现卷积神经网络训练与推理，并使用 tkinter 构建图形化交互界面。系统支持本地视频浏览、播放/暂停、"
        "变速播放、进度拖拽以及 KNN、模板匹配、CNN 三种识别方法的一键切换，能够在视频画面中实时框出数字区域"
        "并标注识别结果。实验表明，OTSU 二值化、居中 padding 和连通域去噪显著改善了推理输入与 MNIST 训练数据"
        "之间的分布差异；CNN 在 MNIST 测试集上达到 98.86% 准确率，并在视频实测中表现出更高的置信度和稳定性。"
    )
    set_font(run)
    para = add_p(doc, "", line=2.0)
    run = para.add_run("关键词：")
    set_font(run, bold=True)
    run = para.add_run("手写数字识别；视频信息处理；卷积神经网络；K-近邻；模板匹配；OpenCV；PyTorch")
    set_font(run)
    doc.add_page_break()

    heading(doc, "一、实验目的")
    for text in [
        "本实验旨在掌握基于机器视觉的模式识别系统构建方法，将视频采集、图像预处理、特征表示、分类器设计和结果可视化连接成完整流程。通过对手写数字视频的实时识别，理解训练数据分布、推理数据预处理和分类模型之间的匹配关系。",
        "具体目标包括：掌握 OpenCV 图像处理与轮廓检测方法；掌握 KNN、模板匹配和 CNN 三类识别算法的基本原理与工程实现；掌握 MNIST 数据集训练、模型保存与加载流程；能够对不同方法的识别效果、置信度、模型大小和实时性进行对比分析；最终形成可交互、可扩展的手写数字视频识别系统。",
    ]:
        add_p(doc, text, first=True)

    heading(doc, "二、系统总体方案")
    add_p(
        doc,
        "系统采用“视频输入层、预处理层、识别层、交互显示层”的分层架构。视频输入层由 VideoProcessor 负责读取本地视频帧并在后台线程中调度识别流程；预处理层由 preprocessing.py 统一完成灰度化、边缘检测、ROI 定位、OTSU 二值化、连通域去噪、居中 padding 和尺寸标准化；识别层提供 KNN、模板匹配和 CNN 三种方法；交互显示层由 tkinter GUI 负责视频显示、方法切换、进度控制和状态反馈。",
        first=True,
    )
    add_placeholder(
        doc,
        "【图1 系统运行主界面截图占位】",
        "建议截图内容：打开视频后的 GUI 主界面，需包含方法选择、播放控制、视频画布和识别框。",
    )
    add_table(
        doc,
        ["模块", "文件", "主要职责"],
        [
            ["入口与参数解析", "src/main.py", "解析模型路径参数并启动 GUI"],
            ["图形界面", "src/gui.py", "构建 tkinter 界面、管理方法切换和状态显示"],
            ["视频处理", "src/video_processor.py", "后台线程读取视频帧并调用识别器"],
            ["共享预处理", "src/preprocessing.py", "ROI 定位、二值化、去噪、居中和尺寸归一化"],
            ["KNN 识别", "src/recognizer.py", "基于 28×28 像素特征和 k=5 进行分类"],
            ["模板匹配", "src/recognizer_template.py", "构造 0-9 平均模板并计算皮尔逊相关系数"],
            ["CNN 识别", "src/recognizer_cnn.py", "加载 PyTorch CNN 模型并输出类别与置信度"],
            ["模型训练", "src/train_cnn.py / src/train_knn.py", "基于 MNIST 训练 CNN 和 KNN 模型"],
        ],
        widths=[3.5, 4.8, 8.2],
    )

    heading(doc, "三、关键方法与实现")
    heading(doc, "3.1 视频帧预处理与 ROI 定位")
    add_p(
        doc,
        "视频帧首先转换为灰度图，再通过形态学膨胀、腐蚀与 absdiff 提取前景差异，结合 Sobel X/Y 边缘响应得到二值边缘图。随后使用轮廓检测筛选候选 ROI，并依据宽度、高度等几何条件过滤明显噪声区域。对每个 ROI，系统进一步执行 OTSU 自适应二值化、连通域去噪、形态学闭运算、可选膨胀、裁剪居中和 resize，使推理输入尽量接近 MNIST 训练样本的“白字黑底、居中、有边距”的分布。",
        first=True,
    )
    add_placeholder(
        doc,
        "【图2 预处理流程截图占位】",
        "建议截图内容：原始 ROI、OTSU 二值化、去噪后图像、居中归一化后 28×28 图像的对比。",
    )

    heading(doc, "3.2 KNN 机器学习方法")
    add_p(
        doc,
        "KNN 方法将标准化后的 28×28 图像展开为 784 维向量，使用 OpenCV KNearest 分类器并设置 k=5，通过欧氏距离寻找最近邻样本并多数投票得到数字类别。初始实现使用 20×20 输入和全部 60,000 个样本，模型体积较大且对噪声敏感；改进后使用 28×28 输入、20,000 个训练样本，并对 KNN 单独禁用膨胀操作，避免像素级距离被笔画加粗带来的系统偏差放大。",
        first=True,
    )

    heading(doc, "3.3 模板匹配方法")
    add_p(
        doc,
        "模板匹配方法从 digits.png 中为 0-9 每类数字生成平均模板，识别时计算 ROI 与各模板之间的皮尔逊相关系数，并选择相关性最高的数字作为输出。该方法不需要训练过程，模型体积极小，具有较强可解释性，但平均模板会抹平同一类别内部的书写差异，因此对真实视频中笔画形变和噪声的适应能力有限。",
        first=True,
    )

    heading(doc, "3.4 CNN 深度学习方法")
    add_p(
        doc,
        "CNN 方法采用两层卷积神经网络：Conv1(1→10, 5×5) + ReLU + MaxPool，Conv2(10→20, 5×5) + ReLU + MaxPool，随后接全连接层 FC1(320→50) 和 FC2(50→10)。模型使用 MNIST 60,000 张训练图像训练 10 个 epoch，在 10,000 张测试集上达到 98.86% 的准确率，模型文件约 89 KB。相比 KNN，CNN 通过卷积核学习局部结构和层次化视觉特征，对笔画粗细、轻微平移和噪声具有更好的鲁棒性。",
        first=True,
    )
    add_placeholder(
        doc,
        "【图3 CNN 训练曲线图片占位】",
        "建议插入 model/cnn_training_curve.png，展示训练损失/准确率变化。",
    )

    heading(doc, "四、问题分析与改进过程")
    heading(doc, "4.1 边缘线条与 MNIST 填充数字的分布不匹配")
    add_p(
        doc,
        "实验初期直接将 Sobel 边缘图作为识别输入，数字只剩细线轮廓，白色像素数量约 100；而 MNIST 训练样本是填充笔画，白色像素通常约 150-200。二者存在明显分布偏移，导致 CNN 置信度一度只有约 0.17。改进策略不是更换模型，而是重构 ROI 标准化流程：使用 OTSU 自适应阈值获得填充数字，使用居中 padding 保持宽高比并模拟 MNIST 外边距，再 resize 到 28×28。该改进将 CNN 置信度提升到约 0.85-1.00，是影响最大的单项优化。",
        first=True,
    )
    add_table(
        doc,
        ["指标", "旧方法：边缘图", "新方法：OTSU+居中", "提升"],
        [
            ["预处理后白像素", "约 100", "约 150-200", "+50%-100%"],
            ["CNN 置信度", "约 0.17", "0.85-1.00", "显著提升"],
            ["KNN-CNN 一致性", "42%", "58%", "+16 个百分点"],
        ],
        widths=[4, 4, 4, 4],
    )

    heading(doc, "4.2 KNN 与 CNN 的预处理参数差异")
    add_p(
        doc,
        "统一预处理后 CNN 明显改善，但 KNN 在部分样本上反而退化。分析发现，CNN 对膨胀操作较鲁棒，而 KNN 使用像素级 L2 距离，膨胀会改变大量像素位置，使测试样本与训练样本的距离关系被扰动。因此系统在共享 prepare_roi() 的同时，为不同算法设置不同参数：CNN 启用 dilate=True 以增强笔画，KNN 使用 dilate=False 以保留原始细节。该经验说明，共享预处理函数并不意味着所有算法应使用完全相同的参数。",
        first=True,
    )
    add_table(
        doc,
        ["参数", "CNN", "KNN 改进后", "原因"],
        [
            ["target_size", "28×28", "28×28", "保留比 20×20 更多的区分性细节"],
            ["dilate", "True", "False", "KNN 对像素级形态变化更敏感"],
            ["训练样本", "60,000", "20,000", "减少离群样本干扰并降低模型体积"],
            ["输出", "类别+Softmax 置信度", "类别", "OpenCV KNN 默认不提供稳定置信度"],
        ],
        widths=[3.5, 3.2, 3.8, 5.5],
    )

    heading(doc, "4.3 连通域去噪")
    add_p(
        doc,
        "手写视频中常出现划痕、试笔点、纸面斑点和拖尾，这些噪声经二值化后与数字主体同为白色。系统利用几何差异进行过滤：数字主体通常是面积最大的连通域，而噪声多为面积较小的零散区域。_remove_small_blobs() 通过 connectedComponentsWithStats 计算各连通域面积，仅保留面积不低于最大连通域 25% 的区域。该方法无需额外训练，计算开销约 1 ms，但在 noise.jpg 样本上将 CNN 置信度从 0.56 提升到 0.70。",
        first=True,
    )
    add_placeholder(
        doc,
        "【图4 噪声样本去噪前后对比图占位】",
        "建议截图内容：noise.jpg 原图、去噪前二值图、去噪后二值图、识别结果。",
    )
    add_table(
        doc,
        ["指标", "去噪前", "去噪后", "变化"],
        [
            ["CNN 置信度", "0.56", "0.70", "+25%"],
            ["白像素数", "37", "29", "划痕被清除"],
            ["预测结果", "5", "5", "正确不变"],
            ["处理时间", "约 0.8 ms", "约 1.2 ms", "开销可忽略"],
        ],
        widths=[4, 4, 4, 4],
    )

    heading(doc, "4.4 三种方法统一管理")
    add_p(
        doc,
        "三种识别方法的模型文件、预处理参数和输出格式均不同。为避免 GUI 中堆积大量 if-else 分支，系统采用工厂模式和统一接口：RecognizerFactory 根据方法名称创建识别器，并将实例缓存；所有识别器实现 preprocess_frame、find_digit_rois、predict_digit、process_frame 和 draw_detections 等一致方法；GUI 只依赖统一接口和 METHODS 元数据。启动时默认加载 CNN，其余模型由后台线程预加载，实现较快启动和无缝切换。",
        first=True,
    )
    add_placeholder(
        doc,
        "【图5 三种方法切换效果截图占位】",
        "建议分别截取 KNN、模板匹配、CNN 三种方法在同一视频帧上的识别画面。",
    )

    heading(doc, "五、实验结果与分析")
    heading(doc, "5.1 训练与模型结果")
    add_p(
        doc,
        "CNN 使用 SGD 优化器训练，学习率 0.01，batch size 为 64，训练 10 个 epoch 后在 MNIST 测试集上准确率达到 98.86%。模型参数量约 21,750，文件大小约 89 KB。KNN 使用 20,000 个 MNIST 样本构建模型，文件约 60 MB；模板匹配只保存 10 个平均模板，体积可忽略。",
        first=True,
    )
    add_table(
        doc,
        ["方法", "训练数据", "模型大小", "置信度", "特点"],
        [
            ["KNN", "MNIST 20k", "约 60 MB", "无", "实现简单，可解释，但泛化能力有限"],
            ["模板匹配", "digits.png 5k", "可忽略", "有但偏低", "零训练成本，受形变影响明显"],
            ["CNN", "MNIST 60k", "约 89 KB", "有且可靠", "准确率最高，实时性和鲁棒性最好"],
        ],
        widths=[3, 3.5, 3, 3, 5],
    )

    heading(doc, "5.2 视频实测结果")
    add_p(
        doc,
        "在测试视频“手写数字.mp4”的同一帧上进行对比，CNN 通常能够以 0.89-1.00 的高置信度输出稳定结果；KNN 可以检测并给出类别，但缺少置信度且连续帧中存在跳变；模板匹配在部分样本中的相关系数仅约 0.08，结果可信度较低。因此本系统推荐以 CNN 作为主识别方法，KNN 和模板匹配作为对比实验方法。",
        first=True,
    )
    add_table(
        doc,
        ["方法", "示例检测数字", "置信度", "评价"],
        [
            ["KNN", "1", "不支持", "可完成分类，但稳定性一般"],
            ["模板匹配", "4", "0.08", "置信度低，可靠性不足"],
            ["CNN", "7", "0.89", "置信度高，结果可信"],
        ],
        widths=[3.5, 4, 3.5, 5],
    )
    add_placeholder(
        doc,
        "【图6 视频实测结果截图占位】",
        "建议截图内容：Frame 800 或 Frame 1200 的识别结果，可优先放 CNN 结果。",
    )

    heading(doc, "5.3 消融分析")
    add_p(
        doc,
        "为了验证各改进项的有效性，实验对预处理步骤进行消融对比。结果表明，OTSU 二值化与居中 padding 是最关键的改进；28×28 分辨率和 KNN 专用参数提升了传统方法的稳定性；连通域去噪进一步增强了复杂手写场景下的鲁棒性。",
        first=True,
    )
    add_table(
        doc,
        ["配置", "CNN 置信度", "KNN-CNN 一致性", "KNN 模型大小"],
        [
            ["基线：edges、20×20、无去噪", "约 0.17", "42%", "92 MB"],
            ["+ OTSU 二值化", "约 0.55", "48%", "92 MB"],
            ["+ 居中 padding", "约 0.80", "52%", "92 MB"],
            ["+ KNN 升级到 28×28", "约 0.85", "55%", "60 MB"],
            ["+ KNN 禁用膨胀", "约 0.88", "58%", "60 MB"],
            ["+ 连通域去噪（最终）", "约 0.90", "58%", "60 MB"],
        ],
        widths=[6, 3.2, 3.5, 3.2],
    )

    heading(doc, "六、结论与展望")
    add_p(
        doc,
        "本实验完成了一个面向手写数字视频的实时模式识别系统，实现了从视频读取、ROI 定位、图像标准化、三类分类器识别到 GUI 可视化展示的完整流程。实验表明，数据预处理对识别效果具有决定性影响；当推理数据被转换为接近 MNIST 的分布后，CNN 置信度由接近随机的 0.17 提升至约 0.90。CNN 在准确率、模型体积和置信度可靠性方面均优于 KNN 与模板匹配，是本任务的最佳方案。",
        first=True,
    )
    add_p(
        doc,
        "后续可从以下方面继续改进：其一，引入目标跟踪或时序投票，利用连续帧信息减少单帧误判；其二，采集真实视频帧并进行标注，用于微调 CNN 以缩小 MNIST 与真实手写视频之间的 domain gap；其三，将当前视频文件输入扩展为摄像头实时输入；其四，使用 PyInstaller 打包为独立可执行程序，降低部署门槛。",
        first=True,
    )

    heading(doc, "参考文献")
    refs = [
        "[1] Y. LeCun, L. Bottou, Y. Bengio, and P. Haffner. Gradient-Based Learning Applied to Document Recognition. Proceedings of the IEEE, 1998, 86(11): 2278-2324.",
        "[2] R. O. Duda, P. E. Hart, and G. D. Stork. Pattern Classification. 2nd Edition. Wiley, 2001.",
        "[3] G. Bradski and A. Kaehler. Learning OpenCV: Computer Vision with the OpenCV Library. O'Reilly Media, 2008.",
        "[4] OpenCV Documentation. K-Nearest Neighbors and Image Processing Modules. https://docs.opencv.org/",
        "[5] PyTorch Documentation. Training a Classifier and Neural Network Modules. https://pytorch.org/tutorials/",
        "[6] MNIST Database of Handwritten Digits. http://yann.lecun.com/exdb/mnist/",
        "[7] CSDN 博客. 手写数字识别——基于 OpenCV KNN. https://blog.csdn.net/reason125132/article/details/124741701",
        "[8] CSDN 博客. PyTorch MNIST CNN 训练. https://blog.csdn.net/qq_45588019/article/details/120935828",
        "[9] 项目 README.md：手写数字视频识别系统使用说明。",
        "[10] 项目 docs/IMPLEMENTATION.md：手写数字视频识别实现文档。",
    ]
    for ref in refs:
        add_p(doc, ref, size=10.5, line=1.5)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing = 1.5
                    for run in para.runs:
                        if run.font.size is None:
                            set_font(run, size=10.5)

    doc.save(str(OUT))


if __name__ == "__main__":
    build()
    print(f"saved {OUT}")
